"""Deterministic checks for round-two reviewer-revision artifacts.

The public checks only require public synthetic aggregate outputs. A separate opt-in
private check inspects the anonymous manuscript without allowing a public clone to
discover adjacent private material.
"""
from __future__ import annotations

import math
import os
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
AUDIT = ROOT.parent
MAIN = ROOT / "outputs" / "final-run-round2-pooled"
ROUND2 = ROOT / "outputs" / "round2-pooled"
BIG4 = ROOT / "outputs" / "round2-big4-pooled"
DOCX = AUDIT / "esg-audit-monte-carlo-private" / "manuscript" / "esg-audit-investment-efficiency-round4-revised-anonymous.docx"


def binary_mcse(p: float, repetitions: int) -> float:
    return math.sqrt(p * (1.0 - p) / repetitions)


def main() -> None:
    main_table = pd.read_csv(MAIN / "tables" / "table-6-independent-seed-crosscheck.csv")
    primary_null = main_table.loc[(main_table["scenario"] == "Null") & (main_table["firms"] == 300)].iloc[0]
    assert int(primary_null["repetitions"]) == 2000
    assert math.isclose(primary_null["firm_rejection_5pct_pooled_mcse"], binary_mcse(primary_null["firm_rejection_5pct"], 2000), rel_tol=1e-12)
    assert math.isclose(primary_null["two_way_rejection_5pct_pooled_mcse"], binary_mcse(primary_null["two_way_rejection_5pct"], 2000), rel_tol=1e-12)
    assert (MAIN / "figures" / "figure-1-primary-operating-characteristics-pooled.png").is_file()

    time = pd.read_csv(ROUND2 / "tables" / "table-10-time-structure-sensitivity-pooled.csv")
    assert len(time) == 12
    assert sorted(time["analysis_time_clusters"].unique().tolist()) == [8, 28]
    assert sorted(time["esg_persistence"].unique().tolist()) == [0.25, 0.6, 0.95]
    assert sorted(time["esg_variance_scale"].unique().tolist()) == [0.0, 1.0]
    assert set(time["repetitions"]) == {1000}
    for _, row in time.iterrows():
        assert math.isclose(row["firm_mcse"], binary_mcse(row["firm_rejection_5pct"], 1000), rel_tol=1e-12)
        assert math.isclose(row["two_way_mcse"], binary_mcse(row["two_way_rejection_5pct"], 1000), rel_tol=1e-12)

    scale = pd.read_csv(ROUND2 / "tables" / "table-11-scale-mapping-pooled.csv")
    assert len(scale) == 6
    assert set(scale["repetitions"]) == {1000}
    assert set(scale["independent_master_seeds"]) == {2}
    assert (scale.loc[scale["effect_scale"] == 0, "gamma_inter_log_sd"] == 0).all()
    assert (scale["oracle_minus_gamma"].abs() < 0.03).all()
    assert (scale["estimated_minus_oracle"].abs() < 0.03).all()

    availability = pd.read_csv(ROUND2 / "tables" / "table-12-selective-availability-pooled.csv")
    assert availability["availability_code"].tolist() == ["complete", "adverse_selective", "coverage_aligned"]
    assert not availability["availability_scenario"].str.contains("MAR", case=False, na=False).any()
    assert set(availability["repetitions"]) == {1000}
    assert availability.loc[availability["availability_code"] == "complete", "mean_second_stage_n"].item() == 2400.0

    fe = pd.read_csv(ROUND2 / "tables" / "table-13-first-stage-fe-sensitivity-pooled.csv")
    assert fe["first_stage_fe"].tolist() == ["industry_year", "industry_plus_year", "none"]
    assert set(fe["repetitions"]) == {1000}
    assert {"mean_oracle_beta_interaction", "oracle_firm_rejection_5pct", "oracle_two_way_rejection_5pct"}.issubset(fe.columns)
    assert (fe["mean_second_stage_n"] == 2400.0).all()
    reference = pd.read_csv(ROUND2 / "tables" / "table-a3-first-stage-fe-design-reference-panels.csv")
    assert set(reference["master_seed"]) == {20260827, 20260828}
    assert (reference["n_first_stage"] == 3000).all()

    big4 = pd.read_csv(BIG4 / "tables" / "table-9-big4-mechanism-ablation-pooled.csv").sort_values("big4_variance_scale")
    assert big4["big4_variance_scale"].tolist() == [0.0, 1.0]
    assert big4["repetitions"].tolist() == [1000, 1000]
    selection_only = big4.iloc[0]
    retained = big4.iloc[1]
    assert abs(selection_only["mean_interaction"]) < 0.01
    assert retained["firm_rejection_5pct"] > selection_only["firm_rejection_5pct"]
    assert (BIG4 / "figures" / "figure-3-big4-mechanism-ablation-pooled.png").is_file()

    for name in ["figure-4-time-structure-sensitivity-pooled.png", "figure-5-scale-mapping-pooled.png", "figure-6-selective-availability-pooled.png"]:
        assert (ROUND2 / "figures" / name).is_file()

    if os.environ.get("VERIFY_PRIVATE_MANUSCRIPT") == "1":
        assert DOCX.is_file(), "Private manuscript audit was requested but no DOCX is present."
        document = Document(DOCX)
        body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        text = body_text + "\n" + table_text
        for required in [
            "Fitted on full panel", "sqrt[p̂(1−p̂)/R]", "unrounded MCSE of 0.00547",
            "Big Four auditor indicator", "coverage-aligned", "[Author name(s) withheld for review]",
            "Table 3a. Paired difference", "No real-company or licensed commercial data are included",
        ]:
            assert required in text, f"Missing required manuscript revision: {required}"
        for forbidden in [
            "MAR-like", "audit-quality proxy", "final-run-calibrated", "manuscript editing assistance",
            "five-company SEC Company Facts metadata snapshot",
        ]:
            assert forbidden not in text, f"Deprecated manuscript wording remains: {forbidden}"
        assert "N=2,000" not in text, "Main effective N remains ambiguous in manuscript text."
        print("Round-two reviewer tests passed, including explicit private anonymous-manuscript audit.")
    else:
        print("Round-two reviewer tests passed; private anonymous-manuscript audit not requested.")


if __name__ == "__main__":
    main()
